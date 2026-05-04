#!/usr/bin/env python
"""
Generate print-ready .docx files for VirtusFocus Focus Group, May 9, 2026.

Produces 8 documents from the locked .txt source content:
- 3 worksheet packets (Institutional, D2C/Athlete, GTM)
- 2 athlete progression visuals (Grace, Angelo)
- 3 artifact packet cover pages (one per segment)

Run:    python make_focus_group_docx.py
Output: .docx files saved alongside their .txt source files.

Functional polish target — clean Word formatting, page breaks, headers
and footers, colored header bands, line-fill writing space, gradient
tier strips on progression visuals.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import subprocess

# LibreOffice path on Windows; used to convert .docx -> .pdf for the print vendor.
SOFFICE = r"C:\Program Files\LibreOffice\program\soffice.exe"


# ============================================================================
# CONFIGURATION
# ============================================================================

ROOT = os.path.dirname(os.path.abspath(__file__))
FOCUS_DIR = os.path.join(ROOT, "Focus Group Materials")

# Segment colors for header bands
SEGMENT_COLORS = {
    "INSTITUTIONAL":   "1F3A60",  # deep blue
    "D2C / ATHLETE":   "2E5E2E",  # forest green
    "GTM":             "C77B30",  # amber
    "ATHLETE":         "2E5E2E",  # green for athlete content (progression visuals)
}

# Tier-strip gradient (light sage -> deep forest), 12 steps
GRADIENT_12 = [
    "C8E6C9", "B7DFB8", "A5D6A7", "93CD96",
    "81C784", "74C078", "66BB6A", "52AB57",
    "388E3C", "2E7D32", "246F2A", "1B5E20",
]

# Angelo's 10-week tier strip: 7 greens + Wk 8 amber + 2 deep greens
GRADIENT_10 = [
    "C8E6C9", "A5D6A7", "81C784", "66BB6A",
    "388E3C", "2E7D32", "246F2A",
    "FFB74D",  # Wk 8 amber (Mixed week)
    "1B5E20", "144418",
]

WRITING_LINE = "_" * 70


# ============================================================================
# OXML HELPERS
# ============================================================================

def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), color_hex)
    tc_pr.append(shading)


def remove_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def set_cell_left_border(cell, color_hex, size=24):
    """Add a colored left border to a cell (for callout effect)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'right', 'bottom']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tc_borders.append(border)
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), str(size))
    left.set(qn('w:color'), color_hex)
    tc_borders.append(left)
    tc_pr.append(tc_borders)


def set_table_no_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.find(qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl.insert(0, tbl_pr)
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{border_name}')
        b.set(qn('w:val'), 'nil')
        borders.append(b)
    tbl_pr.append(borders)


# ============================================================================
# PAGE / FOOTER / DOC SETUP
# ============================================================================

def configure_page(section, top_margin=0.7, side_margin=0.85, bottom_margin=0.7):
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(top_margin)
    section.bottom_margin = Inches(bottom_margin)
    section.left_margin = Inches(side_margin)
    section.right_margin = Inches(side_margin)


def configure_footer(section, segment_label, total_pages):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PAGE ")
    r.font.size = Pt(8)
    r.font.name = "Arial"

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    r._r.append(fldChar1)
    r._r.append(instrText)
    r._r.append(fldChar2)

    r2 = p.add_run(f" of {total_pages}  |  {segment_label}  |  CONFIDENTIAL — NDA APPLIES")
    r2.font.size = Pt(8)
    r2.font.name = "Arial"


def new_document():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = "Arial"
    style.font.size = Pt(11)
    return doc


# ============================================================================
# LAYOUT HELPERS
# ============================================================================

def add_colored_band(doc, text, color_hex, font_color="FFFFFF", width_in=6.8):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, color_hex)
    remove_cell_borders(cell)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell.width = Inches(width_in)

    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string(font_color)
    run.font.name = "Arial"
    return table


def add_centered_heading(doc, text, size_pt=18, bold=True, color="000000"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size_pt)
    r.font.name = "Arial"
    r.font.color.rgb = RGBColor.from_string(color)
    return p


def add_body_paragraph(doc, text, size_pt=11, italic=False, bold=False, indent=0.0,
                       alignment=None, space_before=None, space_after=None):
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(size_pt)
    r.font.name = "Arial"
    r.italic = italic
    r.bold = bold
    return p


def add_horizontal_rule(doc, color_hex="888888", size=6):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_borders = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    p_borders.append(bottom)
    p_pr.append(p_borders)
    return p


def add_writing_lines(doc, count=4):
    """Add N writing lines (long underscores) with tight spacing."""
    for _ in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(WRITING_LINE)
        r.font.size = Pt(11)
        r.font.name = "Courier New"


def add_page_break(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_break(WD_BREAK.PAGE)


def add_checkbox_line(doc, options, indent=0.5):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    for i, opt in enumerate(options):
        if i > 0:
            r_sep = p.add_run("       ")
            r_sep.font.size = Pt(11)
        r = p.add_run(f"[ ]  {opt}")
        r.font.size = Pt(11)
        r.font.name = "Arial"


def add_callout_quote(doc, quote_text, accent_color="2E7D32", indent=0.4):
    """Render an athlete quote as a callout: shaded cell with colored left border."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    cell = table.rows[0].cells[0]
    cell.width = Inches(5.5)
    set_cell_shading(cell, "F5F5F5")  # very light gray
    set_cell_left_border(cell, accent_color, size=24)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    para = cell.paragraphs[0]
    para.paragraph_format.left_indent = Inches(0.15)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    r = para.add_run(quote_text)
    r.italic = True
    r.font.size = Pt(10)
    r.font.name = "Arial"

    # Indent table from left margin
    tblPr = table._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)
    indW = OxmlElement('w:tblInd')
    indW.set(qn('w:w'), str(int(indent * 1440)))
    indW.set(qn('w:type'), 'dxa')
    tblPr.append(indW)


def add_tier_strip(doc, weeks, colors, label="WEEKLY TIER"):
    """Add a single-row table representing the weekly tier visualization."""
    add_body_paragraph(doc, label, size_pt=10, bold=True, indent=0.0,
                       space_before=8, space_after=4)

    # Week numbers row
    table = doc.add_table(rows=2, cols=len(weeks))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_no_borders(table)

    # Top row: week numbers
    for i, wk in enumerate(weeks):
        cell = table.rows[0].cells[i]
        cell.width = Inches(6.5 / len(weeks))
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        r = para.add_run(str(wk))
        r.font.size = Pt(8)
        r.font.name = "Arial"
        r.bold = True

    # Bottom row: colored cells
    for i, color in enumerate(colors):
        cell = table.rows[1].cells[i]
        cell.width = Inches(6.5 / len(weeks))
        set_cell_shading(cell, color)
        para = cell.paragraphs[0]
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        # Empty cell — color is the data
        r = para.add_run(" ")
        r.font.size = Pt(11)


# ============================================================================
# WORKSHEET BUILDERS
# ============================================================================

def add_worksheet_cover(doc, segment_label, segment_color, why_text, format_text):
    add_centered_heading(doc, '"VirtusFocus" Focus Group', size_pt=20)
    add_centered_heading(doc, "May 9, 2026", size_pt=14, bold=False, color="555555")
    doc.add_paragraph()
    add_colored_band(doc, f"{segment_label} BREAKOUT PACKET", segment_color)
    doc.add_paragraph()
    add_centered_heading(doc, "CONFIDENTIAL UNDER NDA", size_pt=12, color="888888")
    add_horizontal_rule(doc)
    doc.add_paragraph()

    # Name fields
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    r = p.add_run("Name (optional): ")
    r.font.size = Pt(11); r.font.name = "Arial"; r.bold = True
    r2 = p.add_run("_" * 50)
    r2.font.size = Pt(11); r2.font.name = "Courier New"

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run("Background (one line): ")
    r.font.size = Pt(11); r.font.name = "Arial"; r.bold = True
    r2 = p.add_run("_" * 45)
    r2.font.size = Pt(11); r2.font.name = "Courier New"

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    r2 = p.add_run("_" * 65)
    r2.font.size = Pt(11); r2.font.name = "Courier New"

    doc.add_paragraph()
    add_horizontal_rule(doc)

    add_body_paragraph(doc, "WHY YOU'RE IN THIS BREAKOUT", size_pt=12, bold=True, indent=0.3,
                       space_before=6)
    add_body_paragraph(doc, why_text, size_pt=11, indent=0.3)
    doc.add_paragraph()
    add_body_paragraph(doc, "THE FORMAT", size_pt=12, bold=True, indent=0.3)
    add_body_paragraph(doc, format_text, size_pt=11, indent=0.3)


def add_question_page(doc, q_label, prompt_lines, sub_prompts=None, checkboxes=None):
    add_page_break(doc)
    add_horizontal_rule(doc)
    add_body_paragraph(doc, q_label, size_pt=14, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER,
                       space_before=4, space_after=4)
    add_horizontal_rule(doc)
    doc.add_paragraph()

    for line in prompt_lines:
        add_body_paragraph(doc, line, size_pt=11, indent=0.3)

    if checkboxes:
        doc.add_paragraph()
        add_checkbox_line(doc, checkboxes)

    if sub_prompts:
        for label, count in sub_prompts:
            if label:
                doc.add_paragraph()
                add_body_paragraph(doc, label, size_pt=11, bold=True, indent=0.3,
                                   space_before=4, space_after=4)
            add_writing_lines(doc, count)
    else:
        doc.add_paragraph()
        add_writing_lines(doc, 8)


def add_anything_else_page(doc):
    add_page_break(doc)
    add_horizontal_rule(doc)
    add_body_paragraph(doc, "ANYTHING ELSE", size_pt=14, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER,
                       space_before=4, space_after=4)
    add_horizontal_rule(doc)
    doc.add_paragraph()
    add_body_paragraph(doc, "What did we not ask that we should have?", size_pt=11,
                       indent=0.3, bold=True, space_before=4, space_after=4)
    add_writing_lines(doc, 7)
    doc.add_paragraph()
    add_body_paragraph(doc, "Anything else for the record:", size_pt=11, indent=0.3,
                       bold=True, space_before=8, space_after=4)
    add_writing_lines(doc, 7)
    doc.add_paragraph()
    add_body_paragraph(doc, "Thank you.", size_pt=11, indent=0.3, italic=True,
                       space_before=8)


# ============================================================================
# THREE WORKSHEET DOCUMENTS
# ============================================================================

def make_institutional_worksheet():
    doc = new_document()
    section = doc.sections[0]
    configure_page(section)
    configure_footer(section, "INSTITUTIONAL", 7)

    why = ("You're being asked these questions because of your experience running "
           "programs and coaching at the next level. We are not looking for "
           "endorsement. We're looking for what we are missing, what would not "
           "work, and what we are getting wrong. Write as much or as little as "
           "you want. The note-taker will capture the discussion, but your "
           "written notes give us signal we would otherwise miss.")
    fmt = ("Five questions, ~9 minutes each. One question per page. A final "
           "page for anything we did not ask but should have. The artifact "
           "you'll review at Question 4 is real output from a real athlete "
           "(anonymized). One of two athletes featured today.")
    add_worksheet_cover(doc, "INSTITUTIONAL", SEGMENT_COLORS["INSTITUTIONAL"], why, fmt)

    add_question_page(doc, "QUESTION 1 — PAST STATE",
        prompt_lines=["When you think about your athletes' mental performance — confidence "
                      "under pressure, recovery from setbacks, focus, identity through "
                      "adversity — what was your program doing about it?"],
        sub_prompts=[("", 4), ("What worked?", 4), ("What did not?", 4)])

    add_question_page(doc, "QUESTION 2 — FAILURE-MOMENT EXTRACTION",
        prompt_lines=["Walk us through the last time something went wrong with an athlete "
                      "on the mental side — pressure response, slump, locker-room thing, "
                      "motivation collapse."],
        sub_prompts=[("The situation:", 4), ("What you actually did:", 4),
                     ("What you wish you could have done:", 4)])

    add_question_page(doc, "QUESTION 3 — STRESS-TEST OUR ASSUMPTION",
        prompt_lines=["Our working assumption: most programs spinning up mental performance "
                      "solutions land on sporadic meetings, on-staff coaches, or guest "
                      "speakers — and those resources don't translate into anything "
                      "athletes use day-to-day."],
        checkboxes=["True", "Untrue", "Partially true"],
        sub_prompts=[("Why:", 6),
                     ("What programs you respect are actually doing that works:", 6)])

    add_question_page(doc, "QUESTION 4 — REACT TO ATHLETE A",
        prompt_lines=["You've reviewed Athlete A's coaching insights, her Wk7 "
                      "failure-recovery output, and her Wk12 peak output."],
        sub_prompts=[("What does this 12-week arc tell you that your coaching tools didn't?", 7),
                     ("What would you want to know about her that this doesn't show you?", 6)])

    add_question_page(doc, "QUESTION 5 — KILLER TEST",
        prompt_lines=[],
        sub_prompts=[("(a)  How would something like this make you win more — concretely?", 6),
                     ("(b)  What is the ONE piece of this that would have to be undeniable "
                      "before you would take it to your AD for budget approval?", 6)])

    add_anything_else_page(doc)

    out_path = os.path.join(FOCUS_DIR, "Worksheets", "Institutional_Worksheet.docx")
    doc.save(out_path)
    return out_path


def make_d2c_worksheet():
    doc = new_document()
    section = doc.sections[0]
    configure_page(section)
    configure_footer(section, "D2C / ATHLETE", 7)

    why = ("You're being asked these questions because you've competed at a high "
           "level, raised athletes who have, or both. We need to know what would "
           "have actually helped you or your athlete — and what would have felt "
           "like another thing to ignore. Honest answers, not polite ones. The "
           "note-taker will capture the discussion; your written notes give us "
           "signal we'd otherwise miss.")
    fmt = ("Five questions, ~9 minutes each. One question per page. A final "
           "page for anything we did not ask but should have. The artifact "
           "you'll review at Question 4 is real output from a real athlete "
           "(anonymized) — one of two athletes featured today.")
    add_worksheet_cover(doc, "D2C / ATHLETE", SEGMENT_COLORS["D2C / ATHLETE"], why, fmt)

    add_question_page(doc, "QUESTION 1 — MEMORY ANCHOR",
        prompt_lines=["Think back to your competitive years — collegiate or high "
                      "school. What's the thing your coaches and your program never "
                      "gave you on the mental side that you needed?",
                      "Confidence stuff, pressure stuff, recovery from mistakes, "
                      "identity stuff, the quiet stuff."],
        sub_prompts=[("What you needed:", 5),
                     ("What got in the way of you getting it:", 5)])

    add_question_page(doc, "QUESTION 2 — ADOPTION-BARRIER EXTRACTION",
        prompt_lines=["If a tool like this had existed during your career — daily "
                      "mindset prompts, weekly coaching created from your own data, "
                      "parent or coach integration (your data and input stays "
                      "private) — would you have actually used it? Honestly."],
        checkboxes=["Yes", "No", "Depends"],
        sub_prompts=[("Why:", 4),
                     ("What would have made it sticky?", 5),
                     ("What would have made it just another thing on your phone you ignore?", 5)])

    add_question_page(doc, "QUESTION 3 — STRESS-TEST OUR ASSUMPTION",
        prompt_lines=["Our working assumption: today's athletes are either getting "
                      "nothing structured around mental performance, or they're "
                      "getting talks and meetings that don't stick."],
        checkboxes=["True", "Untrue", "Partially true"],
        sub_prompts=[("From your kids, your athletes, your network — what have you actually seen?", 11)])

    add_question_page(doc, "QUESTION 4 — REACT TO GRACE'S ARC",
        prompt_lines=["You've reviewed Grace's coaching arc — the messages and deep "
                      "dives she received over twelve weeks."],
        sub_prompts=[("If this were your daughter, your athlete, your younger self — would you want this?", 4),
                     ("What's the strongest part?", 4),
                     ("What's the part that would feel less authentic?", 4)])

    add_question_page(doc, "QUESTION 5 — KILLER TEST",
        prompt_lines=[],
        sub_prompts=[("(a)  If you paid for this directly, what's the monthly price ceiling? "
                      "Where does it stop being worth it?  $ ____ / month", 5),
                     ("(b)  What's the one thing that would make you cancel inside of a month?", 6)])

    add_anything_else_page(doc)

    out_path = os.path.join(FOCUS_DIR, "Worksheets", "D2C_Athlete_Worksheet.docx")
    doc.save(out_path)
    return out_path


def make_gtm_worksheet():
    doc = new_document()
    section = doc.sections[0]
    configure_page(section)
    configure_footer(section, "GTM", 7)

    why = ("You're here because of your experience in sales, marketing, content, "
           "and consumer brand strategy. We need you to challenge the story, the "
           "channel, the architecture, and the price. We are not looking for "
           "endorsement. We're looking for what we are missing, what would not "
           "work, and what we are getting wrong. Write as much or as little as "
           "you want. The note-taker will capture the discussion; your written "
           "notes give us signal we would otherwise miss.")
    fmt = ("Five questions, ~9 minutes each. One question per page. A final "
           "page for anything we did not ask but should have. The artifact "
           "you'll review at Question 4 includes real coaching output and "
           "dashboard mockups.")
    add_worksheet_cover(doc, "GTM", SEGMENT_COLORS["GTM"], why, fmt)

    add_question_page(doc, "QUESTION 1 — PITCH STRESS-TEST",
        prompt_lines=["Forget what we've shown you. In one sentence, how would you "
                      "describe VirtusFocus to someone who's never heard of it?"],
        sub_prompts=[("", 3),
                     ("Where does the story land?", 4),
                     ("Where does it stumble?", 5)])

    add_question_page(doc, "QUESTION 2 — CATEGORY AND COMPETITION",
        prompt_lines=["What category does this compete in, in your head?"],
        checkboxes=["Mental performance coaching", "Athlete development tech",
                    "Wellness app", "Edtech / training platform"],
        sub_prompts=[("Other:", 1),
                     ("Who are we up against?", 4),
                     ("How do we differentiate fast — in one line a parent or coach would actually repeat?", 4)])

    add_question_page(doc, "QUESTION 3 — SINGLE APP vs WHITE-LABEL",
        prompt_lines=["We're betting both channels are viable: institutional (programs, "
                      "athletic departments) AND D2C (parents and athletes directly). "
                      "Different message, different price, different pitch.",
                      "Two architectures to weigh:",
                      "(a) The same app serves both markets — one product, two pitches.",
                      "(b) White-label version institutions can duplicate and \"skin\" "
                      "with their own brand. Ohio State gets \"Ohio State Mental "
                      "Performance\" with their colors, fonts, branding — part of "
                      "their university family."],
        checkboxes=["(a) Same app", "(b) White-label", "Both"],
        sub_prompts=[("Why:", 5),
                     ("What does each cost us that we're not seeing?", 5)])

    add_question_page(doc, "QUESTION 4 — REACT TO THE ARTIFACT",
        prompt_lines=["You've reviewed Grace's coaching output and the dashboard "
                      "mockups."],
        sub_prompts=[("If you were marketing this tomorrow, what's the headline asset?", 4),
                     ("What story actually sells it?", 4),
                     ("What proof are we missing?", 4)])

    add_question_page(doc, "QUESTION 5 — KILLER TEST",
        prompt_lines=["You're advising us on GTM strategy."],
        checkboxes=["Direct-to-consumer first", "Institutional first",
                    "Independent (one path that goes its own way)", "Simultaneously"],
        sub_prompts=[("What does the first $1M of revenue look like?", 5),
                     ("What would you absolutely NOT do?", 5)])

    add_anything_else_page(doc)

    out_path = os.path.join(FOCUS_DIR, "Worksheets", "GTM_Worksheet.docx")
    doc.save(out_path)
    return out_path


# ============================================================================
# PROGRESSION VISUAL BUILDERS
# ============================================================================

def add_visual_header(doc, athlete_label, sport_label, arc_label, dates):
    add_centered_heading(doc, athlete_label.upper(), size_pt=16, bold=True)
    add_centered_heading(doc, f"{sport_label}  |  {arc_label}  |  {dates}",
                         size_pt=10, bold=False, color="555555")
    add_colored_band(doc, "CONFIDENTIAL — NDA APPLIES", "888888", font_color="FFFFFF")


def add_inflection_point(doc, week_label, body_text, quote_text=None):
    """Add one inflection point: bold week label, body paragraph, optional callout quote."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.0)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(week_label + "  ")
    r.bold = True
    r.font.size = Pt(10)
    r.font.name = "Arial"
    r2 = p.add_run(body_text)
    r2.font.size = Pt(10)
    r2.font.name = "Arial"
    if quote_text:
        add_callout_quote(doc, quote_text, accent_color="2E7D32", indent=0.5)


def make_grace_progression():
    doc = new_document()
    section = doc.sections[0]
    configure_page(section, top_margin=0.4, bottom_margin=0.4, side_margin=0.6)
    configure_footer(section, "ATHLETE A", 1)

    add_visual_header(doc,
        "Athlete A — D2 women's softball, outfielder",
        "12-Week Coaching Arc",
        "Athlete A — Progression Visual",
        "Jan 12 – Apr 26, 2026")

    add_body_paragraph(doc, "PHASE TRAJECTORY", size_pt=10, bold=True,
                       space_before=4, space_after=2)
    add_body_paragraph(doc, "Developing  →  Consistent  (advancement at Wk 5)",
                       size_pt=11, indent=0.2, space_after=4)

    add_tier_strip(doc, list(range(1, 13)), GRADIENT_12, "WEEKLY TIER")
    add_body_paragraph(doc,
        "Growth across all 12 weeks. 82 of 84 days = Win. Gradient = building momentum across the arc.",
        size_pt=9, italic=True, indent=0.0, space_before=2, space_after=8)

    add_body_paragraph(doc, "KEY INFLECTION POINTS", size_pt=10, bold=True,
                       space_before=4, space_after=2)

    add_inflection_point(doc, "Wk 1-2",
        "Outcome-reactive baseline at intake. Preparation-energized as program starts.")

    add_inflection_point(doc, "Wk 3",
        "HIGH EI. Derailer activated, structure held. Brief disappointment → reset same-week. Hinge habit operating as designed.")

    add_inflection_point(doc, "Wk 4",
        "Action-proven. First RQS=4 in program (deepest reflection quality available). One competitive chance → breakthrough at-bat:",
        quote_text='"the one chance I got, I did something with."')

    add_inflection_point(doc, "Wk 5",
        "↑ PHASE ADVANCEMENT to Consistent. Identity-embedded statement in journal:",
        quote_text='"If I\'m given a chance, use it well."')

    add_inflection_point(doc, "Wk 6",
        "First competitive validation: pinch-hit single, first collegiate start. NC composure held.")

    add_inflection_point(doc, "Wk 7",
        "FIRST COMPETITIVE FAILURE — pinch-hit strikeout. Reset language by Wednesday. Failure-resilient.")

    add_inflection_point(doc, "Wk 8",
        "Agency-driven. First self-directed development initiative — defensive practice with Kabine on her own time. Attention-management awareness:",
        quote_text='"separate softball from my life."')

    add_inflection_point(doc, "Wk 9",
        "First direct coach advocacy. Second competitive failure → same-day composure.")

    add_inflection_point(doc, "Wk 10",
        "Acceptance-oriented. Zero playing time.",
        quote_text='"I think I\'m just at the point where I don\'t even care to play anymore... I just want everyone to do well."')

    add_inflection_point(doc, "Wk 11",
        "Re-engaged inside acceptance. First Neutral days in 70-day Win streak (challenge-writing, not behavioral disengagement). First specific physical controllable named:",
        quote_text='"running on toes, not backs of feet."')

    add_inflection_point(doc, "Wk 12",
        "ARTICULATED PARADOX COMPOSURE — peak moment. Wednesday breakthrough:",
        quote_text='"fast asf today tbh, got all my mf bunts down."')
    add_callout_quote(doc, '"I did so good, and it still just like doesnt matter. but i cant say it doesnt matter because like it does, its just annoying but once again oh well."',
                      accent_color="2E7D32", indent=0.5)

    doc.add_paragraph()
    add_horizontal_rule(doc)
    add_body_paragraph(doc, "INTAKE → CURRENT", size_pt=10, bold=True,
                       space_before=4, space_after=2)
    add_body_paragraph(doc,
        "FROM: outcome-dependent confidence.  TO: outcome-INDEPENDENT, self-authored recognition.",
        size_pt=11, italic=True, indent=0.2)
    add_body_paragraph(doc,
        "The same athlete who entered the program needing external validation to feel confident is now naming her own performance and sustaining composure when external recognition does not arrive.",
        size_pt=10, indent=0.2, space_before=4)

    out_path = os.path.join(FOCUS_DIR, "Athlete Materials", "Grace Kindel",
                            "Grace_Kindel_12-Week_Progression_Visual.docx")
    doc.save(out_path)
    return out_path


def make_angelo_progression():
    doc = new_document()
    section = doc.sections[0]
    configure_page(section, top_margin=0.4, bottom_margin=0.4, side_margin=0.6)
    configure_footer(section, "ATHLETE B", 1)

    add_visual_header(doc,
        "Athlete B — D2 men's track and field throws",
        "10-Week Coaching Arc",
        "Athlete B — Progression Visual",
        "Feb 2 – Apr 26, 2026")

    add_body_paragraph(doc, "PHASE TRAJECTORY", size_pt=10, bold=True,
                       space_before=4, space_after=2)
    add_body_paragraph(doc,
        "Developing  →  Consistent  (advancement at Wk 10 — final week of athlete's collegiate career after 16 years in T&F)",
        size_pt=11, indent=0.2, space_after=4)

    add_tier_strip(doc, list(range(1, 11)), GRADIENT_10, "WEEKLY TIER (1 Mixed week, 9 Growth)")
    add_body_paragraph(doc,
        "9 of 10 weeks Growth. Wk 8 = first Mixed week of program (Mt Union setback under severe weather + facility loss). 60 Win / 9 Neutral / 2 Missed across 70 program days.",
        size_pt=9, italic=True, indent=0.0, space_before=2, space_after=8)

    add_body_paragraph(doc, "KEY INFLECTION POINTS", size_pt=10, bold=True,
                       space_before=4, space_after=2)

    add_inflection_point(doc, "Wk 1",
        "Illness recovery + approach-framed re-engagement. Returns to competition. Category-only WTD scoring.")

    add_inflection_point(doc, "Wk 2",
        "Near-PR Shotput (16.44m vs 16.70m career best) at highest-level meet of season. Solo at the meet — composure held in isolation.")

    add_inflection_point(doc, "Wk 3",
        "First Gold Streak. Won both events. Identity-level achievement.")

    add_inflection_point(doc, "Wk 4",
        "PSAC championship retrospective after 2.5-week gap. MODERATE EI. Partially Aligned. Multi-domain overload derailer activated under championship pressure. Outdoor transition begins.")

    add_inflection_point(doc, "Wk 5",
        "SHIPPENSBURG BREAKTHROUGH — best outdoor opener of career. All PSAC auto-qualifiers hit. 7/7 Win — second Gold Streak. PSAC lesson named in athlete's own words:",
        quote_text='"properly planned, not spontaneous."')

    add_inflection_point(doc, "Wk 6",
        "First Missed day in 42 program days (Friday illness). CMU meet cancelled. Only 1 practice in 7 days. Moderate recommitment (lifted Saturday while sick). Reframed the disruption:",
        quote_text='"a blessing"')

    add_inflection_point(doc, "Wk 7",
        "Easter break recovery week executed to stated intention. 7/7 Win — third Gold Streak. Bullseye center-ring expanded 1 → 6 items. Achievement internal:",
        quote_text='"mental stability and equilibrium."')

    add_inflection_point(doc, "Wk 8",
        "FIRST MIXED WEEK in program. Mt Union 2-day meet under severe weather. Marks below standard across all three events. Practice facility permanently lost. Strongest negative language to date — but weekday execution held, athlete competed through.")

    add_inflection_point(doc, "Wk 9",
        "SLIPPERY ROCK PROGRAM-BEST. Outdoor shot put school record (over 16m). New Discus PR. 7/7 Win — fourth Gold Streak. Preparation-quality system VALIDATED A SECOND TIME, this time under MORE constrained conditions. Athlete's own summary:",
        quote_text='"best meet of the season."')

    add_inflection_point(doc, "Wk 10",
        "↑ PHASE ADVANCEMENT to Consistent — final week of collegiate career. Two real-time strategic decisions named in athlete's own voice:",
        quote_text='"I didn\'t need to be there but I chose to to support my team and to keep my competition readiness up."')
    add_callout_quote(doc,
        '"leave colligate sports with gratitude, reverence, and peace."',
        accent_color="2E7D32", indent=0.5)

    doc.add_paragraph()
    add_horizontal_rule(doc)
    add_body_paragraph(doc, "INTAKE → CURRENT", size_pt=10, bold=True,
                       space_before=4, space_after=2)
    add_body_paragraph(doc,
        "FROM: multi-domain overload derailer.  TO: preparation-system mastery across three timescales — retrospective lesson extraction (Wk 4), prospective design (Wks 5, 9), in-competition real-time decisions (Wk 10).",
        size_pt=9, italic=True, indent=0.2)

    out_path = os.path.join(FOCUS_DIR, "Athlete Materials", "Angelo Allen",
                            "Angelo_Allen_10-Week_Progression_Visual.docx")
    doc.save(out_path)
    return out_path


# ============================================================================
# ARTIFACT PACKET COVER BUILDERS
# ============================================================================

def add_cover_header(doc, segment_label, segment_color):
    add_centered_heading(doc, '"VirtusFocus" Focus Group', size_pt=20)
    add_centered_heading(doc, "May 9, 2026", size_pt=14, bold=False, color="555555")
    doc.add_paragraph()
    add_colored_band(doc, f"{segment_label} ARTIFACT PACKET", segment_color)
    doc.add_paragraph()
    add_centered_heading(doc, "CONFIDENTIAL UNDER NDA", size_pt=12, color="888888")
    add_horizontal_rule(doc)


def add_cover_section(doc, title, paragraphs, indent=0.3):
    add_body_paragraph(doc, title, size_pt=12, bold=True, indent=indent,
                       space_before=6, space_after=2)
    for p_text in paragraphs:
        add_body_paragraph(doc, p_text, size_pt=11, indent=indent, space_after=2)


def add_cover_athletes(doc, entries, indent=0.3):
    """entries: list of (athlete_label, description) tuples"""
    add_body_paragraph(doc, "ATHLETES FEATURED", size_pt=12, bold=True, indent=indent,
                       space_before=6, space_after=2)
    for label, desc in entries:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(indent)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(11)
        r.font.name = "Arial"
        if desc:
            r2 = p.add_run("  —  " + desc)
            r2.font.size = Pt(11)
            r2.font.name = "Arial"


def add_cover_contents(doc, items, indent=0.3, item_size_pt=11, sub_size_pt=10):
    """items: list of (title, sub_description) tuples"""
    add_body_paragraph(doc, "CONTENTS", size_pt=12, bold=True, indent=indent,
                       space_before=6, space_after=2)
    for i, (title, sub) in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(indent + 0.2)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"{i}.  ")
        r.font.size = Pt(item_size_pt); r.font.name = "Arial"; r.bold = True
        r2 = p.add_run(title)
        r2.font.size = Pt(item_size_pt); r2.font.name = "Arial"
        if sub:
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Inches(indent + 0.7)
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after = Pt(0)
            r3 = p2.add_run(sub)
            r3.italic = True
            r3.font.size = Pt(sub_size_pt); r3.font.name = "Arial"


def add_cover_nda_closing(doc, indent=0.3, with_rule=True):
    if with_rule:
        add_horizontal_rule(doc)
    add_body_paragraph(doc,
        "All material confidential under NDA. Printed copies are collected at end of session; "
        "digital copies are available on request and delivered as personalized watermarked PDFs.",
        size_pt=9, italic=True, indent=indent, space_before=4)


def make_institutional_cover():
    doc = new_document()
    section = doc.sections[0]
    configure_page(section)
    configure_footer(section, "INSTITUTIONAL", 1)

    add_cover_header(doc, "INSTITUTIONAL", SEGMENT_COLORS["INSTITUTIONAL"])

    add_cover_section(doc, "WHAT'S IN THIS PACKET", [
        "Real coaching-staff output from two anonymized athletes, the system's "
        "baseline snapshot showing one athlete's Day 1 starting point, and a "
        "sample institutional team-level dashboard. These are not mockups. This "
        "is what a coaching staff actually receives in real deployment.",
        "Reference these during Question 4 of your breakout."
    ])

    add_cover_athletes(doc, [
        ("Athlete A", "D2 women's softball, outfielder. 12-week coaching arc."),
        ("Athlete B", "D2 men's track and field, throws. 10-week coaching arc, retiring after PSAC championship."),
    ])

    add_cover_contents(doc, [
        ("Athlete A — Baseline Snapshot", "Day 1 starting point"),
        ("Athlete A — 12-week progression visual", None),
        ("Athlete A — Week 7 Coach Insight", "first competitive failure week"),
        ("Athlete A — Week 12 Coach Insight", "peak / phase-sustained"),
        ("Athlete A — Week 12 Coaching Message", "athlete-facing cross-audience reference"),
        ("Athlete A — Week 12 Parent Coaching Message", "parent-facing cross-audience reference"),
        ("Athlete B — 10-week progression visual", None),
        ("Athlete B — Week 10 Coach Insight", "final week + phase advancement"),
        ("Sample Team Snapshot", "institutional aggregate dashboard"),
    ])

    add_cover_nda_closing(doc)

    out_path = os.path.join(FOCUS_DIR, "Artifact Packets", "Institutional_Cover.docx")
    doc.save(out_path)
    return out_path


def make_d2c_cover():
    doc = new_document()
    section = doc.sections[0]
    configure_page(section)
    configure_footer(section, "D2C / ATHLETE", 1)

    add_cover_header(doc, "D2C / ATHLETE", SEGMENT_COLORS["D2C / ATHLETE"])

    add_cover_section(doc, "WHAT'S IN THIS PACKET", [
        "Real athlete-facing coaching output from three anonymized athletes, "
        "plus the system's baseline snapshot for one athlete (Day 1 starting "
        "point) and a sample team dashboard. These are the messages, deep dives, "
        "and parent-facing notes the program actually delivers. Not mockups, "
        "not summaries.",
        "Reference these during Question 4 of your breakout."
    ])

    add_cover_athletes(doc, [
        ("Athlete A", "D2 women's softball, outfielder. 12-week coaching arc."),
        ("Athlete B", "D2 men's track and field, throws. 10-week coaching arc, retiring after PSAC championship."),
        ("Athlete C", "High school wrestling. Week 1 of return arc following season-ending knee injury."),
    ])

    add_cover_contents(doc, [
        ("Athlete A — Baseline Snapshot", "Day 1 starting point"),
        ("Athlete A — 12-week progression visual", None),
        ("Athlete A — Week 7 Coaching Message", "first competitive failure"),
        ("Athlete A — Week 12 Coaching Message", "peak moment"),
        ("Athlete A — Week 12 Deep Dive", "full analytical depth example"),
        ("Athlete A — Week 12 Coach Insight", "coach-facing cross-audience reference"),
        ("Athlete A — Week 12 Parent Coaching Message", "parent-facing perspective"),
        ("Athlete B — 10-week progression visual", None),
        ("Athlete B — Week 10 Coaching Message", "final week + phase advancement"),
        ("Athlete C — Week 1 Parent Coaching Message", "HS-context parent exemplar"),
        ("Sample Team Snapshot", "institutional aggregate dashboard"),
    ], item_size_pt=10, sub_size_pt=9)

    add_cover_nda_closing(doc)

    out_path = os.path.join(FOCUS_DIR, "Artifact Packets", "D2C_Athlete_Cover.docx")
    doc.save(out_path)
    return out_path


def make_gtm_cover():
    doc = new_document()
    section = doc.sections[0]
    configure_page(section, top_margin=0.6, bottom_margin=0.6, side_margin=0.85)
    configure_footer(section, "GTM", 1)

    add_cover_header(doc, "GTM", SEGMENT_COLORS["GTM"])

    add_cover_section(doc, "WHAT'S IN THIS PACKET", [
        "The full product picture. Every audience layer represented. What the "
        "athlete reads, what the coach reads, what the parent reads, what the "
        "coaching staff receives at the team level. Includes baseline snapshots "
        "for all three athletes (Day 1 starting points) and selected app mockup "
        "printouts.",
        "Use these to react to category, story, channel, and architecture during "
        "Question 4 of your breakout."
    ])

    add_cover_athletes(doc, [
        ("Athlete A", "D2 women's softball, outfielder. 12-week coaching arc (depth case)."),
        ("Athlete B", "D2 men's track and field, throws. 10-week coaching arc, retiring after PSAC championship."),
        ("Athlete C", "High school wrestling. Week 1 of return arc post-injury (HS-context parent exemplar)."),
    ])

    add_cover_contents(doc, [
        ("Athlete A — Baseline Snapshot", "Day 1 starting point"),
        ("Athlete A — 12-week progression visual", None),
        ("Athlete A — Week 12 Coaching Message", "athlete-facing example"),
        ("Athlete A — Week 12 Coach Insight", "coach-facing example"),
        ("Athlete A — Week 12 Parent Coaching Message", "parent-facing example"),
        ("Athlete A — Week 12 Parent Insight", "parent dashboard sample"),
        ("Athlete B — Baseline Snapshot", "Day 1 starting point"),
        ("Athlete B — 10-week progression visual", None),
        ("Athlete B — Week 10 Coaching Message", "second athlete example"),
        ("Athlete C — Baseline Snapshot", "Day 1 starting point"),
        ("Athlete C — Week 1 Parent Coaching Message", "HS-context exemplar"),
        ("Sample Team Snapshot", "institutional aggregate dashboard"),
        ("Selected App Mockup Printouts", "Home / Tune-Up / Weekly Recap / Inbox"),
    ], item_size_pt=10, sub_size_pt=9)

    add_cover_nda_closing(doc, with_rule=False)

    out_path = os.path.join(FOCUS_DIR, "Artifact Packets", "GTM_Cover.docx")
    doc.save(out_path)
    return out_path


# ============================================================================
# MAIN
# ============================================================================

if __name__ == "__main__":
    print("=" * 60)
    print('"VirtusFocus" Focus Group — DOCX GENERATOR')
    print("=" * 60)

    os.makedirs(os.path.join(FOCUS_DIR, "Worksheets"), exist_ok=True)
    os.makedirs(os.path.join(FOCUS_DIR, "Artifact Packets"), exist_ok=True)
    os.makedirs(os.path.join(FOCUS_DIR, "Athlete Materials", "Grace Kindel"), exist_ok=True)
    os.makedirs(os.path.join(FOCUS_DIR, "Athlete Materials", "Angelo Allen"), exist_ok=True)

    builders = [
        ("Institutional Worksheet", make_institutional_worksheet),
        ("D2C / Athlete Worksheet",  make_d2c_worksheet),
        ("GTM Worksheet",            make_gtm_worksheet),
        ("Grace Progression Visual", make_grace_progression),
        ("Angelo Progression Visual", make_angelo_progression),
        ("Institutional Cover",      make_institutional_cover),
        ("D2C / Athlete Cover",      make_d2c_cover),
        ("GTM Cover",                make_gtm_cover),
    ]

    docx_paths = []
    for name, fn in builders:
        try:
            path = fn()
            size = os.path.getsize(path)
            rel = os.path.relpath(path, ROOT)
            print(f"  OK   {name:34s}  {size:>6,d} bytes   {rel}")
            docx_paths.append(path)
        except Exception as e:
            print(f"  FAIL {name}: {e}")
            raise

    # Convert all .docx to PDF for the print vendor.
    # Group by output directory so LibreOffice can batch each directory in one call.
    print()
    print("=" * 60)
    print("Converting .docx to .pdf via LibreOffice...")
    print("=" * 60)

    if not os.path.exists(SOFFICE):
        print(f"  WARNING: LibreOffice not found at {SOFFICE}")
        print(f"  Skipping PDF conversion. Install LibreOffice to enable.")
    else:
        by_dir = {}
        for p in docx_paths:
            by_dir.setdefault(os.path.dirname(p), []).append(p)

        for out_dir, files in by_dir.items():
            cmd = [SOFFICE, "--headless", "--convert-to", "pdf",
                   "--outdir", out_dir] + files
            subprocess.run(cmd, capture_output=True, text=True)

        # Verify PDFs were created
        for docx_path in docx_paths:
            pdf_path = docx_path.replace(".docx", ".pdf")
            if os.path.exists(pdf_path):
                size = os.path.getsize(pdf_path)
                rel = os.path.relpath(pdf_path, ROOT)
                print(f"  PDF  {os.path.basename(pdf_path):44s}  {size:>7,d} bytes   {rel}")
            else:
                print(f"  FAIL PDF not created for {docx_path}")

    print()
    print(f"Done. {len(builders)} .docx + matching .pdf files generated.")
