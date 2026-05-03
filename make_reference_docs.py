#!/usr/bin/env python
"""
Convert the three Session 121 reference .txt files into formatted .docx
plus .pdf files for ease of reading and printing.

Inputs (text in worktree):
  Focus Group Materials/Pitch Script/Pitch_Script_v1.txt
  Focus Group Materials/Day-of Operations/Capture_Protocol.txt
  Focus Group Materials/Day-of Operations/Welcome_and_NDA_Flow.txt

Outputs (alongside each .txt):
  *.docx  - formatted Word document
  *.pdf   - LibreOffice headless conversion

Run: python make_reference_docs.py

Formatting strategy (single converter handles all three):
  - First all-caps line  -> Title (large, centered)
  - "====" / "----" rules around a heading line -> Heading 1 (page break above)
  - All-caps line on its own -> Heading 2
  - Lines starting with "*" -> bulleted list item
  - Lines starting with "  - " -> bulleted (sub-level)
  - Lines starting with "[" and ending with "]" -> stage direction (italic, gray)
  - Lines starting with quote marks (full-line "verbatim") -> indented quote block
  - Lines with "Key: value" pattern at line start -> bold key
  - Indented blocks (4+ spaces) -> monospace
  - Default -> body paragraph
"""

import os
import re
import subprocess

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


SOFFICE = r"C:\Program Files\LibreOffice\program\soffice.exe"

ROOT = os.path.dirname(os.path.abspath(__file__))
FOCUS_DIR = os.path.join(ROOT, "Focus Group Materials")

INPUTS = [
    os.path.join(FOCUS_DIR, "Pitch Script", "Pitch_Script_v1.txt"),
    os.path.join(FOCUS_DIR, "Mockup Walkthrough", "Mockup_Walkthrough_Script_v1.txt"),
    os.path.join(FOCUS_DIR, "Day-of Operations", "Capture_Protocol.txt"),
    os.path.join(FOCUS_DIR, "Day-of Operations", "Welcome_and_NDA_Flow.txt"),
    os.path.join(FOCUS_DIR, "Day-of Operations", "Breakout_Facilitation_Guide.txt"),
    os.path.join(FOCUS_DIR, "Day-of Operations", "Note_Taker_Briefing.txt"),
]


# ============================================================================
# DOCX STYLE HELPERS
# ============================================================================

def setup_styles(doc):
    """Configure default styles for body, headings, and special elements."""
    styles = doc.styles

    # Body text
    body = styles["Normal"]
    body.font.name = "Calibri"
    body.font.size = Pt(11)
    body.paragraph_format.space_after = Pt(6)
    body.paragraph_format.line_spacing = 1.15

    # Heading 1 - main section dividers
    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x1F, 0x3A, 0x60)  # deep blue
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)
    h1.paragraph_format.keep_with_next = True

    # Heading 2 - subsections
    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x2E, 0x5E, 0x2E)  # forest green
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(4)
    h2.paragraph_format.keep_with_next = True

    # Heading 3 - sub-subsections
    h3 = styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(2)
    h3.paragraph_format.keep_with_next = True

    # Title
    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(28)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0x1F, 0x3A, 0x60)
    title.paragraph_format.space_after = Pt(12)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def set_page_margins(doc, inches=0.5):
    for section in doc.sections:
        section.top_margin = Inches(inches)
        section.bottom_margin = Inches(inches)
        section.left_margin = Inches(inches)
        section.right_margin = Inches(inches)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_horizontal_rule(doc):
    """Insert a thin horizontal rule via Word's bottom border on an empty para."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "808080")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_stage_direction(doc, text):
    """Italic gray paragraph for [BRACKETED STAGE DIRECTIONS]."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    run.font.size = Pt(10)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)


def add_quote_block(doc, text):
    """Indented blockquote for verbatim athlete language."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.italic = False
    # Subtle left border to mark the quote
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "1F3A60")
    pBdr.append(left)
    pPr.append(pBdr)


def add_monospace(doc, text):
    """Indented monospace paragraph for ASCII tables / code blocks."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.4 * (level + 1))
    _render_inline(p, text)


def add_body(doc, text):
    p = doc.add_paragraph()
    _render_inline(p, text)


def _render_inline(p, text):
    """Render inline markdown-light: **bold**, `code`, plain text."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            p.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith("**"):
            r = p.add_run(token[2:-2])
            r.bold = True
        elif token.startswith("`"):
            r = p.add_run(token[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])


# ============================================================================
# PARSER
# ============================================================================

RE_RULE_EQ = re.compile(r"^=+\s*$")
RE_RULE_DASH = re.compile(r"^-{4,}\s*$")
RE_BULLET = re.compile(r"^(\s*)[\*\-]\s+(.*)$")
RE_NUMBERED = re.compile(r"^(\d+)\.\s+(.*)$")
RE_KEY_VALUE = re.compile(r"^([A-Z][A-Za-z _-]{1,30}):\s+(.*)$")
RE_STAGE_DIR = re.compile(r"^\s*\[[^\]]+\]\s*$")
RE_FULL_QUOTE = re.compile(r'^"[^"]+\.?"\s*$')
RE_SPEAKER_TAG = re.compile(r"^\*\*\[(NICOLE|ARRON)\]\*\*\s*$")
RE_SECTION_HEADING = re.compile(r"^(SECTION|PART)\s+\d+:\s+.+$", re.IGNORECASE)


def is_structural_break(line):
    """Return True if the line breaks paragraph flow (heading, list, rule, etc)."""
    s = line.strip()
    if not s:
        return True
    if RE_RULE_EQ.match(s) or RE_RULE_DASH.match(s):
        return True
    if RE_BULLET.match(line):
        return True
    if RE_STAGE_DIR.match(line):
        return True
    if RE_FULL_QUOTE.match(s):
        return True
    if RE_NUMBERED.match(s) and int(RE_NUMBERED.match(s).group(1)) < 100:
        return True
    if is_all_caps_heading(s) and len(s) > 3:
        return True
    if RE_SECTION_HEADING.match(s):
        return True
    if line.startswith("    ") or line.startswith("\t"):
        return True
    return False


def consume_continuation_lines(lines, start_idx):
    """Read continuation lines after a bullet/numbered/key-value entry.
    Stop at structural breaks. Returns (joined_text, next_idx)."""
    out = []
    j = start_idx
    n = len(lines)
    while j < n:
        nxt = lines[j]
        if is_structural_break(nxt):
            break
        out.append(nxt.strip())
        j += 1
    return (" ".join(out), j)


def is_all_caps_heading(line):
    """All-caps line, 3+ chars, allows numbers/punctuation, no lowercase."""
    s = line.strip()
    if len(s) < 3:
        return False
    if any(c.islower() for c in s):
        return False
    return any(c.isalpha() for c in s)


def parse_and_render(text, doc, mode="generic"):
    """Walk the .txt content and emit docx elements."""
    lines = text.splitlines()
    n = len(lines)
    i = 0

    # Identify the title block: first non-blank line if all-caps
    while i < n and not lines[i].strip():
        i += 1

    # Try to absorb a multi-line title (consecutive all-caps lines + blank)
    title_lines = []
    while i < n and lines[i].strip() and is_all_caps_heading(lines[i]):
        title_lines.append(lines[i].strip())
        i += 1
    if title_lines:
        for tl in title_lines:
            # Skip the rule line if any
            if RE_RULE_EQ.match(tl):
                continue
            p = doc.add_paragraph(tl, style="Title")
        # Skip the leading "====" rule under the title if present
        while i < n and (RE_RULE_EQ.match(lines[i]) or not lines[i].strip()):
            i += 1

    # Walk remaining lines
    in_code = False
    code_buf = []

    def flush_code():
        nonlocal code_buf
        if code_buf:
            add_monospace(doc, "\n".join(code_buf))
            code_buf = []

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Blank line
        if not stripped:
            flush_code()
            i += 1
            continue

        # ==== rule (section divider): page break + the *next* line is a Heading 1
        if RE_RULE_EQ.match(stripped):
            flush_code()
            # Look ahead: heading line + closing ==== rule
            j = i + 1
            heading_line = None
            while j < n and not lines[j].strip():
                j += 1
            if j < n and lines[j].strip() and not RE_RULE_EQ.match(lines[j]):
                heading_line = lines[j].strip()
                # Check for the closing rule
                k = j + 1
                while k < n and not lines[k].strip():
                    k += 1
                if k < n and RE_RULE_EQ.match(lines[k].strip()):
                    add_page_break(doc)
                    doc.add_heading(heading_line, level=1)
                    i = k + 1
                    continue
            # No heading match - just emit a horizontal rule
            add_horizontal_rule(doc)
            i += 1
            continue

        # ---- rule (subtle separator)
        if RE_RULE_DASH.match(stripped):
            flush_code()
            add_horizontal_rule(doc)
            i += 1
            continue

        # SECTION X: heading - upgrade to Heading 1 + page break
        if RE_SECTION_HEADING.match(stripped):
            flush_code()
            add_page_break(doc)
            doc.add_heading(stripped, level=1)
            i += 1
            continue

        # All-caps heading on its own (Heading 2)
        if is_all_caps_heading(stripped) and len(stripped) > 3:
            # Make sure it's not a stage direction (already handled above)
            flush_code()
            doc.add_heading(stripped, level=2)
            i += 1
            continue

        # Speaker tag for pitch script: **[NICOLE]** or **[ARRON]**
        if mode == "pitch" and RE_SPEAKER_TAG.match(stripped):
            flush_code()
            speaker = RE_SPEAKER_TAG.match(stripped).group(1)
            color = (
                RGBColor(0x2E, 0x5E, 0x2E) if speaker == "NICOLE"
                else RGBColor(0xC7, 0x7B, 0x30)
            )
            p = doc.add_paragraph()
            run = p.add_run(f"[{speaker}]")
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = color
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # Stage direction: line is fully [BRACKETED]
        if RE_STAGE_DIR.match(line):
            flush_code()
            add_stage_direction(doc, stripped)
            i += 1
            continue

        # Verbatim quote on its own line (full-line quote)
        if RE_FULL_QUOTE.match(stripped):
            flush_code()
            add_quote_block(doc, stripped)
            i += 1
            continue

        # Bulleted list - consume continuation lines too
        m = RE_BULLET.match(line)
        if m:
            flush_code()
            indent_str = m.group(1)
            content = m.group(2)
            level = 1 if len(indent_str) >= 2 else 0
            cont, next_i = consume_continuation_lines(lines, i + 1)
            if cont:
                content = content + " " + cont
            add_bullet(doc, content, level=level)
            i = next_i
            continue

        # Numbered list (only for short numbers at line start)
        m = RE_NUMBERED.match(stripped)
        if m and int(m.group(1)) < 100:
            flush_code()
            content = m.group(2)
            cont, next_i = consume_continuation_lines(lines, i + 1)
            if cont:
                content = content + " " + cont
            p = doc.add_paragraph(style="List Number")
            _render_inline(p, content)
            i = next_i
            continue

        # Indented monospace block: any line with 2+ leading spaces is treated
        # as part of a structured/preformatted block (preserves ASCII tables,
        # checkbox layouts, indented hierarchies in capture protocol, etc.)
        if line.startswith("  "):
            in_code = True
            code_buf.append(line.rstrip())
            i += 1
            continue
        else:
            if in_code:
                flush_code()
                in_code = False

        # Key: value line (e.g., "Event:        May 9, 2026")
        m = RE_KEY_VALUE.match(line)
        if m and len(m.group(1)) <= 18:
            flush_code()
            value = m.group(2)
            cont, next_i = consume_continuation_lines(lines, i + 1)
            if cont:
                value = value + " " + cont
            p = doc.add_paragraph()
            r = p.add_run(m.group(1) + ": ")
            r.bold = True
            p.add_run(value)
            i = next_i
            continue

        # Default: body paragraph
        flush_code()
        # Group consecutive non-special lines into one paragraph
        para_lines = [stripped]
        i += 1
        while i < n:
            nxt = lines[i]
            nxt_strip = nxt.strip()
            if (not nxt_strip
                or RE_RULE_EQ.match(nxt_strip)
                or RE_RULE_DASH.match(nxt_strip)
                or RE_BULLET.match(nxt)
                or RE_STAGE_DIR.match(nxt)
                or RE_FULL_QUOTE.match(nxt_strip)
                or (mode == "pitch" and RE_SPEAKER_TAG.match(nxt_strip))
                or is_all_caps_heading(nxt_strip)
                or RE_KEY_VALUE.match(nxt)
                or nxt.startswith("  ")
            ):
                break
            para_lines.append(nxt_strip)
            i += 1
        add_body(doc, " ".join(para_lines))

    flush_code()


# ============================================================================
# CONVERSION
# ============================================================================

def convert_one(txt_path):
    docx_path = os.path.splitext(txt_path)[0] + ".docx"
    pdf_dir = os.path.dirname(docx_path)
    pdf_path = os.path.splitext(txt_path)[0] + ".pdf"

    print(f"  Reading:  {os.path.relpath(txt_path, ROOT)}")
    with open(txt_path, "r", encoding="utf-8") as f:
        text = f.read()

    doc = Document()
    setup_styles(doc)
    set_page_margins(doc)  # uses function default; change in set_page_margins() to apply globally

    mode = "pitch" if "Pitch_Script" in os.path.basename(txt_path) else "generic"
    parse_and_render(text, doc, mode=mode)

    doc.save(docx_path)
    print(f"  Wrote:    {os.path.relpath(docx_path, ROOT)}")

    # Convert to PDF via LibreOffice headless
    if os.path.exists(SOFFICE):
        result = subprocess.run(
            [SOFFICE, "--headless", "--convert-to", "pdf", "--outdir",
             pdf_dir, docx_path],
            capture_output=True,
            text=True,
        )
        if os.path.exists(pdf_path):
            print(f"  Wrote:    {os.path.relpath(pdf_path, ROOT)}")
        else:
            print(f"  PDF conversion failed: {result.stderr}")
    else:
        print(f"  LibreOffice not found at {SOFFICE} - skipping PDF")
    print()


def main():
    print("Converting Session 121 reference docs to .docx + .pdf\n")
    for txt_path in INPUTS:
        if not os.path.exists(txt_path):
            print(f"  MISSING:  {txt_path}")
            continue
        convert_one(txt_path)
    print("Done.")


if __name__ == "__main__":
    main()
